# -*- coding: utf-8 -*-
"""mm2docs.py:

## 基本機能
Mindmap (Xmind, FreePlane) から USDM (Excel), Word, Markdown 形式のドキュメントを生成する。

## 利用方法
README.md 参照
"""
import os
import sys
import re
import io
import argparse
import copy
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Tuple, Optional

import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XlImage
from xmind import load as load_xmind_sdk
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage

import mmconfig as cfg
from mmclass import MMNode

class MMConverter:
    """Mindmapデータを読み取り、各種ドキュメントに変換するクラス
    """
    def __init__(self, mm_file: str, template: str, output: str):
        """
        Parameters
        ----------
        mm_file : str
            Source Mindmap file path (.xmind or .mm)
        template : str
            Template file path (.xlsx or .docx)
        output : str
            Output file path (.xlsx, .docx, or .md)
        """
        self.mm_file = mm_file
        self.template = template
        self.output = output
        self.nodes: List[MMNode] = []
        self.doc_data: List[Tuple[str, int]] = [] # Word/MD用 [text, level]

    def load_mm_data(self) -> bool:
        """Mindmapファイルからデータを読み込む
        """
        ext = os.path.splitext(self.mm_file)[1].lower()
        if ext == '.xmind':
            return self._load_xmind()
        elif ext == '.mm':
            return self._load_freeplane()
        else:
            print(f"サポートされていない拡張子です: {ext}")
            return False

    def _load_xmind(self) -> bool:
        """Xmindファイルを読み込む
        """
        try:
            wb = load_xmind_sdk(self.mm_file)
            ws = wb.getPrimarySheet()
            rt = ws.getRootTopic()
            
            raw_nodes = []
            self._get_all_xmind_nodes(rt.getSubTopics(), raw_nodes)
            self._parse_raw_nodes(raw_nodes)
            self._extract_xmind_images()
            return True
        except Exception as e:
            print(f"Xmindファイルの読み込みに失敗しました: {e}")
            return False

    def _get_all_xmind_nodes(self, topics, raw_nodes):
        """Xmindのトピックを再帰的に取得してリストに追加する

        Parameters
        ----------
        topics : list
            XMindのトピックオブジェクトのリスト
        raw_nodes : list
            (テキスト, 画像添付パス or None) のタプルを格納するリスト
        """
        for topic in topics:
            image_src = self._get_topic_image_src(topic)
            raw_nodes.append((topic.getTitle(), image_src))
            if topic.getSubTopics():
                self._get_all_xmind_nodes(topic.getSubTopics(), raw_nodes)

    def _get_topic_image_src(self, topic) -> Optional[str]:
        """トピックのDOM要素からxhtml:img参照を取得する"""
        impl = topic.getImplementation()
        for child in impl.childNodes:
            if hasattr(child, 'tagName') and child.tagName == 'xhtml:img':
                src = child.getAttribute('xhtml:src')
                if src and src.startswith('xap:'):
                    return src[4:]
        return None

    def _load_freeplane(self) -> bool:
        """FreePlane (.mm) ファイルを読み込む
        """
        try:
            tree = ET.parse(self.mm_file)
            root = tree.getroot()

            raw_nodes = []
            # FreePlaneのルートノードの下のノードから取得開始
            for node in root.findall('.//node'):
                text = node.get('TEXT')
                if text:
                    image_src = self._get_freeplane_node_image_src(node)
                    raw_nodes.append((text, image_src))

            self._parse_raw_nodes(raw_nodes)
            self._extract_freeplane_images()
            return True
        except Exception as e:
            print(f"FreePlaneファイルの読み込みに失敗しました: {e}")
            return False

    @staticmethod
    def _get_freeplane_node_image_src(node) -> Optional[str]:
        """FreePlaneノードの直接の子ノードからimg参照を取得する"""
        for child in node.findall('node'):
            for rc in child.findall('richcontent'):
                html_str = ET.tostring(rc, encoding='unicode')
                m = re.search(r'src="([^"]+)"', html_str)
                if m:
                    return m.group(1)
        return None

    def _parse_raw_nodes(self, raw_nodes):
        """生のノード文字列を解析してレベルを決定し、MMNodeリストに格納する

        Parameters
        ----------
        raw_nodes : list
            解析対象のノードテキストのリスト。
            各要素は str または (str, Optional[str]) のタプル。
            タプルの場合、第2要素は画像添付パス。
        """
        stcomment = False
        for item in raw_nodes:
            if isinstance(item, tuple):
                node_text, image_src = item
            else:
                node_text, image_src = item, None

            if not node_text:
                continue
                
            for pattern, level in cfg.PTNLV:
                match = re.match(pattern, node_text)
                if match:
                    # match/case を使用してリファクタリング
                    match level:
                        case cfg.LVNOTEE:
                            if stcomment:
                                stcomment = False
                            break
                        case _ if stcomment:
                            break
                        case cfg.LVNOTE:
                            break
                        case cfg.LVNOTES:
                            if not stcomment:
                                stcomment = True
                            break
                        case l if l <= cfg.LV4:
                            # 見出しレベル1無しで下位が現れたら補完
                            if not self.nodes and l > cfg.LV1:
                                self.nodes.append(MMNode("xxxx xxxx", cfg.LV1))
                            
                            text = match.group(2) if len(match.groups()) >= 2 else node_text
                            
                            # レベル1の場合、キーワード(シート名用)が含まれている可能性があるため
                            # Word/Markdown出力用にキーワードを除去したテキストも保持する
                            if l == cfg.LV1:
                                _, clean_text = self._get_lv1_info(text)
                                self.nodes.append(MMNode(text, l, clean_text))
                            else:
                                nd = MMNode(text, l)
                                if image_src:
                                    nd.image_path = image_src
                                self.nodes.append(nd)
                        case cfg.LVREASON:
                            if self.nodes:
                                text = match.group(2) if len(match.groups()) >= 2 else node_text
                                nd = MMNode(text.replace('\r\n', '\n'), cfg.LVREASON)
                                if image_src:
                                    nd.image_path = image_src
                                self.nodes.append(nd)
                        case cfg.LVREMARK:
                            if self.nodes:
                                text = match.group(2) if len(match.groups()) >= 2 else node_text
                                # 連続する備考は連結
                                if self.nodes[-1].level == cfg.LVREMARK:
                                    self.nodes[-1].text += f"\n{text}"
                                    if image_src and not self.nodes[-1].image_path:
                                        self.nodes[-1].image_path = image_src
                                else:
                                    nd = MMNode(text.replace('\r\n', '\n'), cfg.LVREMARK)
                                    if image_src:
                                        nd.image_path = image_src
                                    self.nodes.append(nd)
                        case cfg.LVSPEC:
                            if self.nodes:
                                nd = MMNode(node_text.replace('\r\n', '\n'), cfg.LVSPEC)
                                if image_src:
                                    nd.image_path = image_src
                                self.nodes.append(nd)
                    break

    def _extract_xmind_images(self):
        """Xmind内の画像をmm_imagesフォルダにPNG形式で保存する

        image_pathにzip内パスが仮格納されているノードを処理し、
        保存後のファイルパスで上書きする。
        """
        image_nodes = [n for n in self.nodes if n.image_path]
        if not image_nodes:
            return

        mm_dir = os.path.dirname(os.path.abspath(self.mm_file))
        image_dir = os.path.join(mm_dir, cfg.IMAGE_DIR)
        os.makedirs(image_dir, exist_ok=True)

        current_kw = "img"
        img_seq = {}

        with zipfile.ZipFile(self.mm_file, 'r') as zf:
            for node in self.nodes:
                if node.level == cfg.LV1:
                    current_kw, _ = self._get_lv1_info(node.text)
                    if current_kw not in img_seq:
                        img_seq[current_kw] = 1

                if not node.image_path:
                    continue

                zip_path = node.image_path
                seq = img_seq.get(current_kw, 1)
                img_seq[current_kw] = seq + 1

                sanitized = self._sanitize_filename(node.text, max_len=50)
                filename = f"{current_kw}_{str(seq).zfill(2)}_{sanitized}.png"
                filepath = os.path.join(image_dir, filename)

                try:
                    image_data = zf.read(zip_path)
                    img = PILImage.open(io.BytesIO(image_data))
                    img.save(filepath, 'PNG')
                    node.image_path = filepath
                    print(f"  画像抽出: {filename}")
                except Exception as e:
                    print(f"  画像抽出エラー ({zip_path}): {e}")
                    node.image_path = ""

    def _extract_freeplane_images(self):
        """FreePlane .mm内の画像をmm_imagesフォルダにPNG形式で保存する

        image_pathに.mmファイルからの相対パスが仮格納されている
        ノードを処理し、保存後のファイルパスで上書きする。
        """
        image_nodes = [n for n in self.nodes if n.image_path]
        if not image_nodes:
            return

        mm_dir = os.path.dirname(os.path.abspath(self.mm_file))
        image_dir = os.path.join(mm_dir, cfg.IMAGE_DIR)
        os.makedirs(image_dir, exist_ok=True)

        current_kw = "img"
        img_seq = {}

        for node in self.nodes:
            if node.level == cfg.LV1:
                current_kw, _ = self._get_lv1_info(node.text)
                if current_kw not in img_seq:
                    img_seq[current_kw] = 1

            if not node.image_path:
                continue

            src_path = os.path.join(mm_dir, node.image_path)
            if not os.path.exists(src_path):
                print(f"  画像ファイルが見つかりません: {src_path}")
                node.image_path = ""
                continue

            seq = img_seq.get(current_kw, 1)
            img_seq[current_kw] = seq + 1

            sanitized = self._sanitize_filename(node.text, max_len=50)
            filename = f"{current_kw}_{str(seq).zfill(2)}_{sanitized}.png"
            filepath = os.path.join(image_dir, filename)

            try:
                img = PILImage.open(src_path)
                img.save(filepath, 'PNG')
                node.image_path = filepath
                print(f"  画像抽出: {filename}")
            except Exception as e:
                print(f"  画像抽出エラー ({src_path}): {e}")
                node.image_path = ""

    @staticmethod
    def _sanitize_filename(text: str, max_len: int = 50) -> str:
        """ファイル名に使えない文字を置換し、長さを制限する"""
        name = text.split('\n')[0]
        name = re.sub(cfg.IMAGE_FILENAME_PROHIBIT, cfg.IMAGE_FILENAME_SUB, name)
        if len(name) > max_len:
            name = name[:max_len]
        return name

    def save(self) -> bool:
        """指定された拡張子に応じて保存処理を行う
        """
        ext = os.path.splitext(self.output)[1].lower()
        match ext:
            case '.xlsx':
                return self.save_xlsx()
            case '.docx':
                return self.save_docx()
            case '.md':
                return self.save_md()
            case _:
                print(f"サポートされていない出力形式です: {ext}")
                return False

    def save_xlsx(self) -> bool:
        """USDM Excel形式で保存
        """
        print(f"Excel (USDM) 出力中: {self.output}")
        if not os.path.splitext(self.template)[1].lower() == '.xlsx':
            print("エラー: テンプレートファイルは .xlsx 形式である必要があります。")
            return False

        try:
            wb = openpyxl.load_workbook(self.template)
            tmpl_sheet = wb[cfg.usdm_cfg.tmpl_sheet]
            
            # 書式情報の取得
            default_styles = []
            for col in range(1, tmpl_sheet.max_column + 1):
                default_styles.append(tmpl_sheet.cell(row=cfg.usdm_cfg.row_start, column=col)._style)

            scnt = 1
            sn_list = []
            
            # データの転記
            current_sheet = None
            rownum = cfg.usdm_cfg.row_start
            id_cnt = [0, 0, 0]
            sn = ""

            for node in self.nodes:
                match node.level:
                    case cfg.LV1:
                        # 新しいシートの作成
                        current_sheet = wb.copy_worksheet(tmpl_sheet)
                        
                        # シート名設定
                        kw, lv1_txt = self._get_lv1_info(node.text)
                        sn = kw
                        if sn in sn_list:
                            for i in range(1, 101):
                                if f"{sn}{i}" not in sn_list:
                                    sn = f"{sn}{i}"
                                    break
                        sn_list.append(sn)
                        
                        current_sheet.title = f"{cfg.usdm_cfg.usdm_sheet}_{sn}"
                        wb.move_sheet(current_sheet, offset=-wb.index(current_sheet) + scnt)
                        scnt += 1
                        
                        # 初期化
                        current_sheet.delete_rows(cfg.usdm_cfg.row_start, current_sheet.max_row + 1)
                        current_sheet[cfg.usdm_cfg.lv1_title_cell].value = lv1_txt
                        id_cnt = [0, 0, 0]
                        rownum = cfg.usdm_cfg.row_start - 1
                        
                        # ウィンドウ枠の固定
                        current_sheet.freeze_panes = cfg.usdm_cfg.freeze_panes

                    case cfg.LVREMARK:
                        if current_sheet and rownum > cfg.usdm_cfg.row_start:
                            # 直前の行に備考をセット
                            # 元の Xmind2USDM.py のロジック: rownum = rownum - 1
                            # そして setcelldata(..., cfg.TXTREASON) を呼び出す
                            rownum -= 1
                            self.setcelldata(current_sheet, rownum, cfg.usdm_cfg.col_remark, node.text, cfg.TXTREASON, default_styles)

                    case cfg.LV2:
                        id_cnt = [id_cnt[0] + 1, 0, 0]
                        data = [""] * 7
                        data[cfg.usdm_cfg.col_req - 1] = cfg.TXTREQ2
                        data[cfg.usdm_cfg.col_req] = f"{sn}-{str(id_cnt[0]).zfill(2)}"
                        data[cfg.usdm_cfg.col_req + 1] = node.text
                        self.setrowdata(current_sheet, rownum, data, cfg.TXTREQ2, default_styles)

                    case cfg.LV3:
                        id_cnt = [id_cnt[0], id_cnt[1] + 1, 0]
                        data = [""] * 7
                        data[cfg.usdm_cfg.col_sub - 1] = cfg.TXTREQ3
                        data[cfg.usdm_cfg.col_sub] = f"{sn}-{str(id_cnt[0]).zfill(2)}-{str(id_cnt[1]).zfill(2)}"
                        data[cfg.usdm_cfg.col_sub + 1] = node.text
                        self.setrowdata(current_sheet, rownum, data, cfg.TXTREQ3, default_styles)

                    case cfg.LV4:
                        data = [""] * 7
                        data[cfg.usdm_cfg.col_grp - 1] = f"＜{node.text}＞"
                        self.setrowdata(current_sheet, rownum, data, cfg.TXTREQ4, default_styles)

                    case cfg.LVREASON:
                        data = [""] * 7
                        data[cfg.usdm_cfg.col_reason - 2] = cfg.TXTREASON
                        data[cfg.usdm_cfg.col_reason - 1] = node.text
                        self.setrowdata(current_sheet, rownum, data, cfg.TXTREASON, default_styles)

                    case cfg.LVSPEC:
                        id_cnt = [id_cnt[0], id_cnt[1], id_cnt[2] + 1]
                        data = [""] * 7
                        data[cfg.usdm_cfg.col_spec - 2] = f"{sn}-{str(id_cnt[0]).zfill(2)}-{str(id_cnt[1]).zfill(2)}-{str(id_cnt[2]*10).zfill(3)}"
                        data[cfg.usdm_cfg.col_spec - 1] = node.text
                        self.setrowdata(current_sheet, rownum, data, cfg.TXTSPEC, default_styles)

                if current_sheet:
                    # 最初のデータ行（4行目）の高さ自動調整修正
                    if rownum == 4:
                        current_sheet.row_dimensions[rownum].height = None
                    rownum += 1

            # 画像を持つノードの別シート作成
            for node in self.nodes:
                if not node.image_path or not os.path.exists(node.image_path):
                    continue
                img_basename = os.path.splitext(os.path.basename(node.image_path))[0]
                sheet_name = re.sub(cfg.usdm_cfg.prohibit_char, cfg.usdm_cfg.sub_txt, img_basename)[:31]
                img_sheet = wb.create_sheet(title=sheet_name)
                title_cell = img_sheet[cfg.usdm_cfg.img_title_cell]
                title_cell.value = sheet_name
                title_cell.font = cfg.usdm_cfg.img_title_font
                xl_img = XlImage(node.image_path)
                img_sheet.add_image(xl_img, cfg.usdm_cfg.img_insert_cell)

            wb.save(self.output)
            return True
        except Exception as e:
            print(f"Excel出力中にエラーが発生しました: {e}")
            return False

    def _get_lv1_info(self, txt: str) -> Tuple[str, str]:
        """大要求レベルのノード情報をキーワードとテキストに分割する

        Parameters
        ----------
        txt : str
            大要求ノードのテキスト

        Returns
        -------
        kw : str
            キーワード（シート名に使用、禁止文字置換済み）
        lv1_txt : str
            大要求の内容テキスト
        """
        try:
            parts = txt.split(None, 1)
            kw_raw = parts[0]
            lv1_txt = parts[1] if len(parts) > 1 else parts[0]
        except:
            kw_raw = lv1_txt = txt
        
        kw = re.sub(cfg.usdm_cfg.prohibit_char, cfg.usdm_cfg.sub_txt, kw_raw)
        return kw, lv1_txt

    def setrowdata(self, sht, rnum, rdata, lv, default_styles):
        """1行分のテキスト、書式をセット
        """
        for cn, rd in enumerate(rdata):
            self.setcelldata(sht, rnum, cn + 1, rd, lv, default_styles)
            # 古い Xmind2USDM.py の挙動を再現：ループ内で毎回呼び出す
            self.set_modulestyle(sht, rnum, default_styles)

    def setcelldata(self, sht, rnum, cnum, txt, lv, default_styles):
        """セルにテキスト、書式をセット
        """
        cell = sht.cell(row=rnum, column=cnum)
        cell.value = txt
        cell._style = copy.deepcopy(default_styles[cnum - 1])
        if lv in cfg.usdm_cfg.cell_style:
            cell.fill = cfg.usdm_cfg.cell_style[lv][cnum - 1][0]
            cell.font = cfg.usdm_cfg.cell_style[lv][cnum - 1][1]

    def set_modulestyle(self, sht, rnum, default_styles):
        """「関連モジュール」列に罫線設定
        """
        for i, cs in enumerate(default_styles[cfg.usdm_cfg.col_module_start-1:]):
            cell = sht.cell(row=rnum, column=cfg.usdm_cfg.col_module_start+i)
            cell._style = copy.deepcopy(cs)

    def save_docx(self) -> bool:
        """Word形式で保存
        """
        print(f"Word出力中: {self.output}")
        if not os.path.splitext(self.template)[1].lower() == '.docx':
            print("エラー: テンプレートファイルは .docx 形式である必要があります。")
            return False

        try:
            doc = Document(self.template)
            for node in self.nodes:
                has_image = node.image_path and os.path.exists(node.image_path)
                if has_image:
                    img_title = os.path.splitext(os.path.basename(node.image_path))[0]

                match node.level:
                    case l if l <= cfg.word_cfg.max_head_lv:
                        doc.add_paragraph('') # 見出しの前に改行
                        doc.add_heading(node.clean_text, level=l + cfg.word_cfg.head_lv_offset)
                        if has_image:
                            doc.add_picture(node.image_path, width=Inches(6))
                            doc.add_paragraph(img_title, style='Caption')
                    case cfg.LVREMARK if has_image:
                        doc.add_picture(node.image_path, width=Inches(6))
                        doc.add_paragraph(img_title, style='Caption')
                    case cfg.LVREMARK:
                        doc.add_paragraph(node.text, style=cfg.word_cfg.txt_remark)
                    case cfg.LVREASON:
                        doc.add_paragraph(node.text, style=cfg.word_cfg.txt_reason)
                    case _ if has_image:
                        doc.add_picture(node.image_path, width=Inches(6))
                        doc.add_paragraph(img_title, style='Caption')
                    case _:
                        doc.add_paragraph(node.text)
            
            doc.save(self.output)
            return True
        except Exception as e:
            print(f"Word出力中にエラーが発生しました: {e}")
            return False

    def save_md(self) -> bool:
        """Markdown形式で保存
        """
        print(f"Markdown出力中: {self.output}")
        try:
            md_text = cfg.md_cfg.header
            for node in self.nodes:
                has_image = node.image_path and os.path.exists(node.image_path)
                text = self._escape_md(node.clean_text)
                # ノード内の改行を <br /> に変換
                text = text.replace('\n', '<br />')

                if has_image:
                    img_title = os.path.splitext(os.path.basename(node.image_path))[0]
                    img_rel = os.path.relpath(node.image_path,
                                              os.path.dirname(os.path.abspath(self.output)))
                    img_rel = img_rel.replace('\\', '/')

                if has_image:
                    img_tag = f'<img src="{img_rel}" alt="{img_title}" width="{cfg.md_cfg.img_width}">\n\n'

                match node.level:
                    case l if l <= cfg.word_cfg.max_head_lv:
                        md_text += f'{"#" * l} {text}\n\n'
                        if has_image:
                            md_text += f'*{img_title}*\n\n'
                            md_text += img_tag
                    case cfg.LVREMARK if has_image:
                        md_text += f'*{img_title}*\n\n'
                        md_text += img_tag
                    case cfg.LVREMARK:
                        md_text += f'{cfg.md_cfg.section_remark}{text}\n{cfg.md_cfg.sec_close}'
                    case cfg.LVREASON:
                        md_text += f'{cfg.md_cfg.section_reason}{text}\n{cfg.md_cfg.sec_close}'
                    case _ if has_image:
                        md_text += f'*{img_title}*\n\n'
                        md_text += img_tag
                    case _:
                        md_text += f'{text}\n\n'
            
            with open(self.output, 'w', encoding='utf-8') as f:
                f.write(md_text)
            return True
        except Exception as e:
            print(f"Markdown出力中にエラーが発生しました: {e}")
            return False

    def _escape_md(self, txt: str) -> str:
        """Markdownの特殊文字をエスケープする
        ただし、バックトークンで囲まれた部分は除外する

        Parameters
        ----------
        txt : str
            エスケープ対象のテキスト

        Returns
        -------
        str
            エスケープ後のテキスト
        """
        # バックトークンで囲まれた部分はエスケープしない
        # それ以外の部分をエスケープ対象文字列で置換する
        parts = re.split(r'(`[^`]*`)', txt)
        for i in range(len(parts)):
            # 偶数番目の要素（バックトークンの外側）のみエスケープ処理を行う
            if i % 2 == 0:
                for s in cfg.md_cfg.escape_strs:
                    parts[i] = parts[i].replace(s, '\\' + s)
        return "".join(parts)

def main():
    parser = argparse.ArgumentParser(description='Mindmap to Documents Converter')
    parser.add_argument('mm', help='Source Mindmap file (.xmind or .mm)')
    parser.add_argument('-t', '--template', help='Template file (.xlsx or .docx)', default='')
    parser.add_argument('-o', '--output', help='Output file (.xlsx, .docx, or .md)', default='')

    args = parser.parse_args()

    # パス解決
    mm_file = args.mm
    if not os.path.exists(mm_file):
        print(f"エラー: 指定されたファイルが見つかりません: {mm_file}")
        sys.exit(1)

    output = args.output
    if not output:
        # デフォルト出力先の設定
        base = os.path.splitext(mm_file)[0]
        output = base + "_USDM.xlsx" # デフォルトはExcel

    template = args.template
    if not template:
        # 出力形式に応じたデフォルトテンプレート
        ext = os.path.splitext(output)[1].lower()
        if ext == '.xlsx':
            template = cfg.usdm_cfg.tmpl_file
        elif ext == '.docx':
            template = cfg.word_cfg.tmpl_name

    converter = MMConverter(mm_file, template, output)
    if converter.load_mm_data():
        if converter.save():
            print("変換完了！")
        else:
            print("変換失敗。")
            sys.exit(1)
    else:
        print("Mindmapの読み込みに失敗しました。")
        sys.exit(1)

if __name__ == '__main__':
    main()
